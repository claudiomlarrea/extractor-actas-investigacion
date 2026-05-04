import streamlit as st
import gspread
from pathlib import Path
from google.oauth2.service_account import Credentials
from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor


def _fila_sheet_normalizada(r):
    return {k.lower().strip(): v for k, v in r.items()}


def _unidad_academica_clave(r):
    """Misma fuente que los encabezados de Word: columna «unidad académica» o «unidad»."""
    row = _fila_sheet_normalizada(r)
    return str(row.get("unidad académica") or row.get("unidad") or "").strip()


def ordenar_registros_por_unidad_academica(registros):
    """Orden estable por unidad; dentro de cada unidad conserva el orden del sheet."""
    if not registros:
        return registros
    indexed = list(enumerate(registros))
    indexed.sort(key=lambda t: (_unidad_academica_clave(t[1]).casefold(), t[0]))
    return [r for _, r in indexed]


# =========================
# ⚙ CONFIGURACIÓN
# =========================

st.set_page_config(page_title="Consejo de Investigación", layout="wide")

_APP_ROOT = Path(__file__).resolve().parent.parent
_LOGO_PATH = _APP_ROOT / "assets" / "logo_uccuyo.png"

# =========================
# 🎨 HEADER
# =========================

col1, col2 = st.columns([1, 8], vertical_alignment="center")

with col1:
    if _LOGO_PATH.is_file():
        st.image(str(_LOGO_PATH), width=120)
    else:
        st.caption("Logo no encontrado (assets/logo_uccuyo.png)")

with col2:
    st.markdown("""
    <div class='header-uccuyo'>
        <h2 style="font-weight:600;">Universidad Católica de Cuyo</h2>
        <h4 style="opacity:0.9;">Secretaría de Investigación</h4>
        <h5 style="opacity:0.8;">Consejo de Investigación</h5>
    </div>
    """, unsafe_allow_html=True)

# =========================
# 🎨 CSS GLOBAL
# =========================

st.markdown("""
<style>

/* Fondo general */
.stApp {
    background-color: #E6E6E6;
}

/* HEADER INSTITUCIONAL */
.header-uccuyo {
    background: linear-gradient(90deg, #064a3f, #0B6B5D);
    padding: 20px;
    border-radius: 10px;
    display: flex;
    flex-direction: column;
    justify-content: center;
    box-shadow: 0px 4px 10px rgba(0,0,0,0.15);
}

.header-uccuyo h2,
.header-uccuyo h4,
.header-uccuyo h5 {
    color: white !important;
    margin: 0;
}

/* Títulos GENERALES (NO TOCA HEADER) */
h1, h2, h3, h4, h5, h6 {
    color: black !important;
}

/* Labels */
label {
    color: black !important;
    font-weight: 700 !important;
}

/* INPUTS */
input, textarea {
    background-color: white !important;
    color: black !important;
}

/* Selectbox */
div[data-baseweb="select"] {
    background-color: white !important;
}

div[data-baseweb="select"] span {
    color: black !important;
}

div[role="listbox"] {
    background-color: white !important;
}

div[role="option"] {
    background-color: white !important;
    color: black !important;
}

div[role="option"]:hover {
    background-color: #e6e6e6 !important;
}

/* Barra vertical al final del texto en selects = caret del combobox Base Web */
div[data-baseweb="select"] input,
div[data-baseweb="select"] [data-baseweb="input"] input,
div[data-baseweb="select"] [role="combobox"],
div[data-baseweb="select"] [contenteditable="true"] {
    caret-color: transparent !important;
}

/* Placeholder */
::placeholder {
    color: #777 !important;
}

.stTextInput > div > div > input {
    background-color: white !important;
    color: black !important;
    caret-color: #111111 !important;
}

.stTextArea textarea {
    background-color: white !important;
    color: black !important;
    caret-color: #111111 !important;
}

/* Números: mismo cursor oscuro sobre fondo blanco */
.stNumberInput input,
[data-testid="stNumberInputField"] {
    caret-color: #111111 !important;
}
/* SIDEBAR OPCIONES EN VERDE */
[data-testid="stSidebarNav"] a {
    background-color: #064a3f !important;
    color: white !important;
    font-size: 16px !important;
    margin-bottom: 8px;
    padding: 10px 12px;
    border-radius: 10px;
}

/* OPCIÓN ACTIVA */
[data-testid="stSidebarNav"] a[aria-current="page"] {
    background-color: #0B6B5D !important;
    color: white !important;
    font-weight: bold;
}

</style>
""", unsafe_allow_html=True)

# =========================
# 🔐 CONEXIÓN
# =========================

scope = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

creds = Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=scope
)

client = gspread.authorize(creds)

SHEET_ID = "17MiyW17W7oLIwSCKjDXCoA85CwBkYqHYhDKblVN37c8"
sheet = client.open_by_key(SHEET_ID).worksheet("Hoja 2")

st.markdown(
    """
    <div style='background-color:#d4edda; padding:15px; border-radius:10px'>
        <a href='https://docs.google.com/spreadsheets/d/17MiyW17W7oLIwSCKjDXCoA85CwBkYqHYhDKblVN37c8' target='_blank' style='text-decoration:none; color:#155724; font-weight:bold; font-size:18px;'>
            ✔ Conectado a Google Sheets (abrir)
        </a>
    </div>
    """,
    unsafe_allow_html=True
)

# =========================
# 📅 DATOS ACTAS
# =========================

actas_dict = {
    187: {"mes": "Febrero"},
    188: {"mes": "Marzo"},
    189: {"mes": "Abril"},
    190: {"mes": "Mayo"},
    191: {"mes": "Junio"},
    192: {"mes": "Julio"},
    193: {"mes": "Agosto"},
    194: {"mes": "Septiembre"},
    195: {"mes": "Octubre"},
    196: {"mes": "Noviembre"},
    197: {"mes": "Diciembre"},
}

fechas_actas = {
    187: "19 de Febrero 2026",
    188: "19 de Marzo 2026",
    189: "16 de Abril 2026",
    190: "21 de Mayo 2026",
    191: "18 de Junio 2026",
    192: "23 de Julio 2026",
    193: "20 de Agosto 2026",
    194: "15 de Septiembre 2026",
    195: "22 de Octubre 2026",
    196: "19 de Noviembre 2026",
    197: "10 de Diciembre 2026"
}
# 🔹 LISTA DE CATEGORÍAS (VA ARRIBA DE TODO)
categoria_opciones = [
    "Seleccionar",
    "Investigador/a Superior I",
    "Investigador/a Principal II",
    "Investigador/a Independiente III",
    "Investigador/a Adjunto/a IV",
    "Investigador/a Asistente V",
    "Becario/a de Iniciación VI",
    "Sin categorización / Externo"
]

opciones_unidades = [
    "Seleccionar unidad académica",
    "FDCSSL- Facultad de Derecho y Ciencias Sociales Sede San Luis",
    "FCMSL- Facultad de Ciencias Médicas Sede San Luis",
    "FCVSL- Facultad de Ciencias Veterinarias Sede San Luis",
    "FCEESL- Facultad de Ciencias Económicas y Empresariales Sede San Luis",
    "FBOSCO- Facultad Don Bosco",
    "FCEESJ- Facultad de Ciencias Económicas y Empresariales Sede San Juan",
    "FFyHSJ- Facultad de Filosofía y Humanidades",
    "ISDSM- Instituto Universitario Santa María",
    "ECRyPSJ- Escuela Cultura Religiosa y Pastoral",
    "FDCSSJ- Facultad de Derecho y Ciencias Sociales Sede San Juan",
    "FCMSJ- Facultad de Ciencias Médicas San Juan",
    "FEDSJ- Facultad de Educación",
    "ESEGSJ- Escuela de Seguridad",
    "FCQyTSJ- Facultad de Ciencias Químicas y Tecnológicas",
    "ISB- Instituto San Buenaventura",
    "Secretaría de Investigación",
    "Unidad de Vinculación Tecnológica",
]

# =========================
# 📝 FORMULARIO
# =========================

st.subheader("Sistema de gestión de temas para el Consejo de Investigación")
st.markdown("<span style='color:black; font-weight:700;'>🔷 Complete solo los campos que correspondan</span>", unsafe_allow_html=True)

st.markdown(
    "<div style='margin-bottom:-15px; color:black; font-weight:700;'>🟢 Elija actividad</div>",
    unsafe_allow_html=True
)
tipo = st.selectbox("", [
        "Proyecto de Investigación",
        "Proyecto de Cátedra",
        "Informe Final",
        "Informe de Avance",
        "Jornada de Investigación",
        "Convocatoria a Proyectos de investigación",
        "Creación de Semillero de Investigación",
        "Categorización Docente",
        "Cronograma",
        "Llamado a Concurso de Becas",
        "Líneas prioritarias de investigación",
        "Otra"
    ])

with st.form("form_acta", clear_on_submit=True):

    # =========================
    # 📅 DATOS BÁSICOS
    # =========================

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Año</div>", unsafe_allow_html=True)
    anio = st.text_input("", "2026", key="anio")

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Número de Acta</div>", unsafe_allow_html=True)

    acta_label = st.selectbox(
        "",
        [f"Acta N°{n} - {actas_dict[n]['mes']}" for n in actas_dict],
        key="acta"
    )

    numero_acta = int(acta_label.split(" ")[1].replace("N°", ""))

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Fecha</div>", unsafe_allow_html=True)
    
    fecha = st.selectbox(
        "",
        list(fechas_actas.values()),
        index=list(fechas_actas.keys()).index(numero_acta),
        key="fecha"
    )

    # =========================
    # 📌 IDENTIFICACIÓN
    # =========================

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Denominación de la actividad</div>", unsafe_allow_html=True)
    titulo = st.text_input("")

    st.markdown("""
    <div style="margin-top:-25px; margin-bottom:5px; background-color:#E6E6E6; padding:10px; border-radius:5px; font-size:13px; color:#000000;">
    <span style="font-weight:700; color:#000000;">Indicaciones:</span>
    <ul style="margin-top:5px; margin-bottom:0; color:#000000;">
    <li style="color:#000000;">Título del proyecto</li>
    <li style="color:#000000;">Informe Final o de Avance</li>
    <li style="color:#000000;">Jornada / semillero / instituto</li>
    </ul>
    </div>
    """, unsafe_allow_html=True)

    # =========================
    # 🎯 PUNTAJE
    # =========================

    puntaje = 0
    if tipo in ["Proyecto de Investigación", "Proyecto de Cátedra", "Informe Final", "Informe de Avance"]:
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Puntaje</div>", unsafe_allow_html=True)
    
        puntaje = st.number_input(
            "",
            min_value=0.0,
            max_value=1000.0,
            step=0.1,
            format="%.2f",
            value=0.0,
            key="puntaje"
        )
    # =========================
    # 🧾 DESCRIPCIÓN
    # =========================

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Descripción</div>", unsafe_allow_html=True)
    descripcion = st.text_area("")

    # =========================
    # 👥 EQUIPO (CONDICIONAL)
    # =========================

    director = ""
    cat_director = ""
    codirector = ""
    categoria_codirector = ""
    equipo = ""
    instituto = ""
    catedra = ""
    alumnos = ""
    apellido_nombre_docente = ""
    dni_docente = ""

    if tipo == "Categorización Docente":
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Apellido y nombre del docente</div>", unsafe_allow_html=True)
        apellido_nombre_docente = st.text_input("", key="apellido_nombre_docente")
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 DNI</div>", unsafe_allow_html=True)
        dni_docente = st.text_input("", key="dni_docente")

    if tipo in ["Proyecto de Investigación", "Proyecto de Cátedra", "Informe Final", "Informe de Avance", "Otra"]:

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Director</div>", unsafe_allow_html=True)
        director = st.text_input("", key="director")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Categoría del Director</div>", unsafe_allow_html=True)
        cat_director = st.selectbox("", categoria_opciones, key="cat_director")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Codirector</div>", unsafe_allow_html=True)
        codirector = st.text_input("", key="codirector")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Categoría del Codirector</div>", unsafe_allow_html=True)
        categoria_codirector = st.selectbox("", categoria_opciones, key="cat_codirector")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Equipo de Investigación</div>", unsafe_allow_html=True)
        equipo = st.text_area("", key="equipo")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Instituto</div>", unsafe_allow_html=True)
        instituto = st.text_input("", key="instituto")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Cátedra</div>", unsafe_allow_html=True)
        catedra = st.text_input("", key="catedra")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Cantidad de Alumnos en el proyecto</div>", unsafe_allow_html=True)
        alumnos = st.text_input("", key="alumnos")

    # =========================
    # 🏫 UNIDAD
    # =========================

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Unidad Académica</div>", unsafe_allow_html=True)
    unidad = st.selectbox("", opciones_unidades, key="unidad")

    # =========================
    # 📄 RESOLUCIONES
    # =========================

    if tipo in ["Proyecto de Investigación", "Proyecto de Cátedra", "Informe Final", "Informe de Avance", "Otra"]:

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Resolución CD</div>", unsafe_allow_html=True)
        resolucion_cd = st.text_input("", key="resolucion_cd")

        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Resolución CS (solo para Informes Finales y de Avances)</div>", unsafe_allow_html=True)
        resolucion_cs = st.text_input("", key="resolucion_cs")

    else:
        resolucion_cd = ""
        resolucion_cs = ""

    # =========================
    # 👤 RESPONSABLE
    # =========================

    st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🔴 Responsable de carga</div>", unsafe_allow_html=True)
    responsable_de_carga = st.text_input("", key="responsable")

    # =========================
    # 💰 FINANCIAMIENTO
    # =========================

    if tipo != "Categorización Docente":
    
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Tipo de financiamiento</div>", unsafe_allow_html=True)
        tipo_financiamiento = st.selectbox("", ["Seleccionar...", "Interno", "Externo"], key="fin")
    
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Fuente</div>", unsafe_allow_html=True)
        fuente_financiamiento = st.text_input("", key="fuente")
    
        st.markdown("<div style='margin-bottom:-10px; color:black; font-weight:700;'>🟢 Monto en pesos (sin puntos)</div>", unsafe_allow_html=True)
        monto_financiamiento = st.number_input("", min_value=0, step=1000, value=None, key="monto")
    
    else:
        tipo_financiamiento = ""
        fuente_financiamiento = ""
        monto_financiamiento = 0

    # =========================
    # 🔘 SUBMIT
    # =========================

    submit = st.form_submit_button("Clic para Enviar al Consejo de Investigación (Google Sheets")
# =========================
# 💾 GUARDAR
# =========================

if submit:

    if tipo_financiamiento == "Seleccionar...":
        tipo_financiamiento = ""

    # 🔹 LIMPIAR "Seleccionar" (versión robusta)
    if unidad.strip().startswith("Seleccionar"):
        unidad = ""

    if instituto.strip().startswith("Seleccionar"):
        instituto = ""

    if catedra.strip().startswith("Seleccionar"):
        catedra = ""

    if tipo_financiamiento.strip().startswith("Seleccionar"):
        tipo_financiamiento = ""

    if str(cat_director).strip().startswith("Seleccionar"):
        cat_director = ""

    if str(categoria_codirector).strip().startswith("Seleccionar"):
        categoria_codirector = ""

    fila = [
        numero_acta,
        fecha,
        anio,
        tipo,
        titulo,
        descripcion,
        director,
        cat_director,
        codirector,
        categoria_codirector,
        equipo,
        apellido_nombre_docente,
        dni_docente,
        unidad,
        resolucion_cd,
        resolucion_cs,
        instituto,
        catedra,
        tipo_financiamiento,
        fuente_financiamiento,
        monto_financiamiento,
        alumnos,
        puntaje,
        responsable_de_carga
        ]

    # VALIDACIONES
    if not anio.strip():
        st.error("Debe completar el año")

    elif not numero_acta:
        st.error("Debe seleccionar el número de acta")

    elif not fecha:
        st.error("Debe seleccionar la fecha")

    elif not tipo:
        st.error("Debe elegir la actividad")

    else:
        sheet.append_row(fila)
        st.success("Registro guardado correctamente")
    
# =========================
# 📄 GENERAR WORD
# =========================

st.markdown("## 📄 Generar Orden del Día")

acta_word = st.selectbox(
    "Seleccionar Orden del Día para generarlo y descargar",
    options=[f"{n} - {actas_dict[n]['mes']}" for n in actas_dict]
)

generar = st.button("Generar Word")

if generar:

    datos = sheet.get_all_records()

    acta_num = int(acta_word.split(" - ")[0])

    registros = [
        r for r in datos
        if str(r.get("numero_acta", "")).strip() == str(acta_num)
    ]
    registros = ordenar_registros_por_unidad_academica(registros)

    if not registros:
        st.warning("No hay registros para esta acta")

    else:
        doc = Document()

        doc.add_heading('Consejo de Investigación', 0)

        p_acta = doc.add_paragraph(f'Acta N° {acta_num}')
        p_acta.paragraph_format.space_after = Pt(0)

        fecha_real = registros[0].get("FECHA", registros[0].get("fecha", ""))
        p_fecha = doc.add_paragraph(f'Fecha: {fecha_real}')
        p_fecha.paragraph_format.space_after = Pt(0)

        doc.add_heading('Orden del Día', level=1)

        contador = 1
        unidad_actual = None

        for r in registros:

            r = {k.lower().strip(): v for k, v in r.items()}

            unidad = r.get("unidad académica", r.get("unidad", "")).strip()

            if unidad != unidad_actual:
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(6)
                h.paragraph_format.space_after = Pt(2)
                h.paragraph_format.line_spacing = 1

                run_h = h.add_run(unidad)
                run_h.bold = True
                run_h.font.color.rgb = RGBColor(0, 102, 204)

                unidad_actual = unidad

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1

            p.add_run(f"{contador}. {r.get('tipo', '')} - {r.get('titulo', '')}\n").bold = True

            descripcion = r.get("descripcion") or r.get("descripción") or ""

            if descripcion:
                p.add_run(f"   Descripción: {descripcion}\n")
                
            tipo_actividad = r.get("tipo", "")

            if tipo_actividad == "Categorización Docente":

                nombre_doc = r.get("apellido_nombre_docente", "")
                dni_doc = r.get("dni_docente", "")

                if nombre_doc:
                    p.add_run(f"   Docente: {nombre_doc}\n")

                if dni_doc:
                    p.add_run(f"   DNI: {dni_doc}\n")

            tipos_con_director = [
                "Proyecto de Investigación",
                "Proyecto de Cátedra",
                "Informe Final",
                "Informe de Avance"
            ]

            if tipo_actividad in tipos_con_director:

                cat = r.get('cat_director', '')
                if cat == "Seleccionar" or cat == "":
                    p.add_run(f"   Director: {r.get('director', '')}\n")
                else:
                    p.add_run(f"   Director: {r.get('director', '')} ({cat})\n")

                cat_codir = r.get('cat_codirector', '')
                if cat_codir == "Seleccionar" or cat_codir == "":
                    p.add_run(f"   Codirector: {r.get('codirector', '')}\n")
                else:
                    p.add_run(f"   Codirector: {r.get('codirector', '')} ({cat_codir})\n")

            equipo_txt = r.get("equipo", "")

            if equipo_txt:
                equipo_txt = equipo_txt.replace("\n", "; ")
                p.add_run(f"   Equipo: {equipo_txt}\n")

            p.add_run(f"   Unidad Académica: {unidad}\n")

            puntaje_valor = r.get("puntaje", 0)
            try:
                puntaje_num = float(puntaje_valor)
            except:
                puntaje_num = 0

            if puntaje_num > 0:
                p.add_run(f"   Puntaje: {int(puntaje_num)}\n")

            if r.get("resolucion_cd"):
                p.add_run(f"   Resolución CD: {r.get('resolucion_cd')}\n")

            if r.get("resolucion_cs"):
                p.add_run(f"   Resolución CS del Proyecto: {r.get('resolucion_cs')}\n")

            if r.get("instituto"):
                p.add_run(f"   Instituto: {r.get('instituto')}\n")

            if r.get("catedra"):
                p.add_run(f"   Cátedra: {r.get('catedra')}\n")

            if r.get("tipo de financiamiento"):
                p.add_run(f"   Financiamiento: {r.get('tipo de financiamiento')}\n")

            if r.get("fuente de financiamiento"):
                p.add_run(f"   Fuente: {r.get('fuente de financiamiento')}\n")

            if r.get("responsable_de_carga"):
                p.add_run(f"   Responsable de carga: {r.get('responsable_de_carga')}\n")

            if r.get("monto del financiamiento"):
                try:
                    monto = int(float(r.get("monto del financiamiento")))
                    monto = f"${monto:,}".replace(",", ".")
                except:
                    monto = r.get("monto del financiamiento")

                p.add_run(f"   Monto: {monto}\n")

            if r.get("alumnos"):
                p.add_run(f"   Alumnos: {r.get('alumnos')}\n")

            contador += 1

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "Descargar Word",
            data=buffer,
            file_name=f"Acta_{acta_num}.docx"
        )
        
st.markdown("### 🧾 Generar informe por responsable")

responsable_reporte = st.text_input("Responsable de carga para generar informe")

generar_responsables = st.button("Generar informe del responsable")

if generar_responsables:

    datos = sheet.get_all_records()

    acta_num = int(acta_word.split(" - ")[0])

    registros = [
        r for r in datos
        if str(r.get("numero_acta", "")).strip() == str(acta_num)
        and str(r.get("responsable_de_carga", "")).strip().lower() == responsable_reporte.strip().lower()
    ]
    registros = ordenar_registros_por_unidad_academica(registros)

    if not responsable_reporte.strip():
        st.warning("Debe ingresar el responsable de carga")

    elif not registros:
        st.warning("No hay registros cargados por ese responsable para esta acta")

    else:
        doc = Document()

        doc.add_heading("Informe de temas cargados", 0)
        doc.add_paragraph(f"Acta N° {acta_num}")
        doc.add_paragraph(f"Responsable de carga: {responsable_reporte}")

        contador = 1
        unidad_actual = None

        for r in registros:

            r = {k.lower().strip(): v for k, v in r.items()}

            unidad = r.get("unidad académica", r.get("unidad", "")).strip()

            if unidad != unidad_actual:
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(6)
                h.paragraph_format.space_after = Pt(2)

                run_h = h.add_run(unidad)
                run_h.bold = True
                run_h.font.color.rgb = RGBColor(0, 102, 204)

                unidad_actual = unidad

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1

            p.add_run(f"{contador}. {r.get('tipo', '')} - {r.get('titulo', '')}\n").bold = True

            descripcion = r.get("descripcion") or r.get("descripción") or ""

            if descripcion:
                p.add_run(f"   Descripción: {descripcion}\n")

            p.add_run(f"   Unidad Académica: {unidad}\n")

            contador += 1

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "Descargar informe del responsable",
            data=buffer,
            file_name=f"Informe_{responsable_reporte}_Acta_{acta_num}.docx"
        )
