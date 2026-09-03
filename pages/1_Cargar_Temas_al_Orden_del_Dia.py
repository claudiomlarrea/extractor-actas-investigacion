from __future__ import annotations

import re
import unicodedata
from html import escape

import streamlit as st
import streamlit.components.v1 as components
import gspread
from gspread.utils import ValueInputOption, InsertDataOption
from pathlib import Path
from ucc_streamlit_chrome import hide_streamlit_cloud_toolbar
from google.oauth2.service_account import Credentials
from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor
import smtplib
import ssl
from email.message import EmailMessage


def _fila_sheet_normalizada(r):
    return {k.lower().strip(): v for k, v in r.items()}


def _unidad_academica_clave(r):
    """Misma fuente que los encabezados de Word: columna «unidad académica» o «unidad»."""
    row = _fila_sheet_normalizada(r)
    return str(row.get("unidad académica") or row.get("unidad") or "").strip()


def _numero_acta_igual(valor, acta_num) -> bool:
    """Compara número de acta aunque Sheets lo devuelva como 193, '193' o 193.0."""
    if valor is None or valor == "":
        return False
    try:
        return int(float(str(valor).strip().replace(",", "."))) == int(acta_num)
    except (TypeError, ValueError):
        return str(valor).strip() == str(acta_num).strip()


def _parse_numero_acta(valor) -> int | None:
    """Convierte numero_acta a int si es posible (193, '193' o 193.0)."""
    if valor is None or valor == "":
        return None
    try:
        return int(float(str(valor).strip().replace(",", ".")))
    except (TypeError, ValueError):
        return None


def _contar_temas_acta(worksheet, acta_num) -> int:
    try:
        datos = worksheet.get_all_records()
    except Exception:
        return 0
    return sum(1 for r in datos if _numero_acta_igual(r.get("numero_acta"), acta_num))


def ordenar_registros_por_unidad_academica(registros):
    """Orden estable por unidad; dentro de cada unidad conserva el orden del sheet."""
    if not registros:
        return registros
    indexed = list(enumerate(registros))
    indexed.sort(key=lambda t: (_unidad_academica_clave(t[1]).casefold(), t[0]))
    return [r for _, r in indexed]


def _fingerprint_tema(r) -> str:
    """Identifica un tema de forma estable para reordenar el Word (no es ID de Sheets)."""
    row = _fila_sheet_normalizada(r)
    return "|".join(
        [
            str(row.get("tipo", "")).strip(),
            str(row.get("titulo", "") or row.get("título", "")).strip(),
            str(row.get("responsable_de_carga", "")).strip(),
            str(row.get("unidad académica") or row.get("unidad") or "").strip(),
            str(row.get("director", "")).strip(),
            str(row.get("descripcion") or row.get("descripción") or "").strip()[:120],
        ]
    )


def _clave_orden_manual_od(acta_num) -> str:
    return f"od_orden_manual_{acta_num}"


def _aplicar_orden_manual_od(acta_num, registros):
    """Aplica el orden elegido por la asistente; temas nuevos van al final de su unidad."""
    if not registros:
        return registros
    key = _clave_orden_manual_od(acta_num)
    by_fp: dict[str, list] = {}
    fps_sheet = []
    for r in registros:
        fp = _fingerprint_tema(r)
        fps_sheet.append(fp)
        by_fp.setdefault(fp, []).append(r)

    prev = st.session_state.get(key)
    if not isinstance(prev, list) or not prev:
        base = ordenar_registros_por_unidad_academica(registros)
        st.session_state[key] = [_fingerprint_tema(r) for r in base]
        return base

    usados: dict[str, int] = {}
    ordenados = []
    for fp in prev:
        cola = by_fp.get(fp) or []
        idx = usados.get(fp, 0)
        if idx < len(cola):
            ordenados.append(cola[idx])
            usados[fp] = idx + 1

    for fp in fps_sheet:
        cola = by_fp.get(fp) or []
        idx = usados.get(fp, 0)
        if idx < len(cola):
            ordenados.append(cola[idx])
            usados[fp] = idx + 1

    st.session_state[key] = [_fingerprint_tema(r) for r in ordenados]
    return ordenados


def _mover_tema_od(acta_num, fingerprint: str, delta: int, registros) -> bool:
    """Mueve un tema ↑/↓ solo dentro de la misma unidad académica."""
    key = _clave_orden_manual_od(acta_num)
    orden = list(st.session_state.get(key) or [])
    if fingerprint not in orden:
        return False
    i = orden.index(fingerprint)
    j = i + delta
    if j < 0 or j >= len(orden):
        return False
    by_fp = {_fingerprint_tema(r): r for r in registros}
    ri = by_fp.get(orden[i])
    rj = by_fp.get(orden[j])
    if ri is None or rj is None:
        return False
    if _unidad_academica_clave(ri).casefold() != _unidad_academica_clave(rj).casefold():
        return False
    orden[i], orden[j] = orden[j], orden[i]
    st.session_state[key] = orden
    st.session_state.pop(f"od_docx_bytes_{acta_num}", None)
    return True


def _ui_reordenar_temas_od(acta_num, registros) -> None:
    """↑/↓ por unidad. No escribe en Google Sheets."""
    if not registros:
        return
    st.markdown("### Reordenar temas del Orden del Día")
    st.caption(
        "Dentro de cada unidad académica, use ↑ o ↓ para definir cómo salen en el Word. "
        "No modifica la planilla de Google Sheets."
    )

    unidad_actual = None
    for idx, r in enumerate(registros):
        row = _fila_sheet_normalizada(r)
        unidad = _unidad_academica_clave(r)
        if unidad != unidad_actual:
            st.markdown(f"**{unidad or 'Sin unidad académica'}**")
            unidad_actual = unidad

        fp = _fingerprint_tema(r)
        tipo = str(row.get("tipo", "")).strip()
        titulo = str(row.get("titulo", "") or row.get("título", "")).strip()
        resp = str(row.get("responsable_de_carga", "")).strip()
        etiqueta = f"{tipo} — {titulo}" if titulo else tipo
        if resp:
            etiqueta = f"{etiqueta} ({resp})"

        c_up, c_dn, c_txt = st.columns([0.08, 0.08, 0.84])
        with c_up:
            if st.button("↑", key=f"od_up_{acta_num}_{idx}", help="Subir dentro de la unidad"):
                if _mover_tema_od(acta_num, fp, -1, registros):
                    st.rerun()
        with c_dn:
            if st.button("↓", key=f"od_dn_{acta_num}_{idx}", help="Bajar dentro de la unidad"):
                if _mover_tema_od(acta_num, fp, 1, registros):
                    st.rerun()
        with c_txt:
            st.markdown(f"{idx + 1}. {etiqueta}")


def _construir_bytes_orden_del_dia(acta_num, registros) -> bytes:
    """Arma el Word del Orden del Día con el orden actual de registros."""
    doc = Document()

    doc.add_heading("Consejo de Investigación", 0)

    p_acta = doc.add_paragraph(f"Acta N° {acta_num}")
    p_acta.paragraph_format.space_after = Pt(0)

    fecha_real = registros[0].get("FECHA", registros[0].get("fecha", ""))
    p_fecha = doc.add_paragraph(f"Fecha: {fecha_real}")
    p_fecha.paragraph_format.space_after = Pt(0)

    doc.add_heading("Orden del Día", level=1)

    contador = 1
    unidad_actual = None

    for r in registros:
        r = {k.lower().strip(): v for k, v in r.items()}
        unidad = r.get("unidad académica", r.get("unidad", "")).strip()

        if unidad != unidad_actual:
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(2)
            h.paragraph_format.line_spacing = 1
            run_h = h.add_run(unidad)
            run_h.bold = True
            run_h.font.color.rgb = RGBColor(0, 102, 204)
            unidad_actual = unidad

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1
        p.add_run(f"{contador}. {r.get('tipo', '')} - {r.get('titulo', '')}\n").bold = True

        descripcion = r.get("descripcion") or r.get("descripción") or ""
        if descripcion:
            p.add_run(f"   Descripción: {descripcion}\n")

        tipo_actividad = r.get("tipo", "")

        if tipo_actividad == "Categorización Docente":
            nombre_doc = r.get("apellido_nombre_docente", "")
            dni_doc = r.get("dni_docente", "")
            if nombre_doc:
                p.add_run(f"   Docente: {nombre_doc}\n")
            if dni_doc:
                p.add_run(f"   DNI: {dni_doc}\n")

        tipos_con_director = [
            "Proyecto de Investigación",
            "Proyecto de Cátedra",
            "Informe Final",
            "Informe de Avance",
        ]
        if tipo_actividad in tipos_con_director:
            cat = r.get("cat_director", "")
            if cat == "Seleccionar" or cat == "":
                p.add_run(f"   Director: {r.get('director', '')}\n")
            else:
                p.add_run(f"   Director: {r.get('director', '')} ({cat})\n")

            cat_codir = r.get("cat_codirector", "")
            if cat_codir == "Seleccionar" or cat_codir == "":
                p.add_run(f"   Codirector: {r.get('codirector', '')}\n")
            else:
                p.add_run(f"   Codirector: {r.get('codirector', '')} ({cat_codir})\n")

        equipo_txt = r.get("equipo", "")
        if equipo_txt:
            equipo_txt = equipo_txt.replace("\n", "; ")
            p.add_run(f"   Equipo: {equipo_txt}\n")

        p.add_run(f"   Unidad Académica: {unidad}\n")

        raw_punt = r.get("puntaje")
        txt_puntaje = puntaje_texto_para_word(raw_punt)
        if txt_puntaje:
            p.add_run(f"   Puntaje: {txt_puntaje}\n")

        if r.get("resolucion_cd"):
            p.add_run(f"   Resolución CD: {r.get('resolucion_cd')}\n")
        if r.get("resolucion_cs"):
            p.add_run(f"   Resolución CS del Proyecto: {r.get('resolucion_cs')}\n")
        if r.get("instituto"):
            p.add_run(f"   Instituto: {r.get('instituto')}\n")

        catedra_od = (
            r.get("catedra")
            or r.get("cátedra")
            or r.get("catedras")
            or r.get("cátedras")
            or ""
        )
        if str(catedra_od).strip():
            etiqueta = (
                "Cátedras"
                if (";" in str(catedra_od) or "," in str(catedra_od))
                else "Cátedra"
            )
            p.add_run(f"   {etiqueta}: {catedra_od}\n")

        if r.get("tipo de financiamiento"):
            p.add_run(f"   Financiamiento: {r.get('tipo de financiamiento')}\n")
        if r.get("fuente de financiamiento"):
            p.add_run(f"   Fuente: {r.get('fuente de financiamiento')}\n")
        if r.get("responsable_de_carga"):
            p.add_run(f"   Responsable de carga: {r.get('responsable_de_carga')}\n")

        if r.get("monto del financiamiento"):
            try:
                monto = int(float(r.get("monto del financiamiento")))
                monto = f"${monto:,}".replace(",", ".")
            except Exception:
                monto = r.get("monto del financiamiento")
            p.add_run(f"   Monto: {monto}\n")

        if r.get("alumnos"):
            p.add_run(f"   Alumnos: {r.get('alumnos')}\n")

        contador += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


TIPOS_CON_PUNTAJE = [
    "Proyecto de Investigación",
    "Proyecto de Cátedra",
    "Informe Final",
    "Informe de Avance",
]


def _ayuda_en_iframe(html: str, alto: int) -> None:
    """HTML en iframe: el CSS del tema de Streamlit no vuelve blanco el texto de ayuda."""
    components.html(html, height=alto, scrolling=False)


def _normalizar_puntaje_desde_hoja(x: float) -> float:
    """Corrige escalas x100 típicas de Google Sheets (columna %, formato numérico). Ej: 8850 → 88.5."""
    if x != x or x <= 0:
        return x
    x = float(x)
    while x > 1000 and abs(x - round(x)) < 1e-4:
        ri = int(round(x))
        if ri % 100 != 0:
            break
        x = x / 100.0
    return x


def parse_puntaje_valor(val):
    """Número desde Sheets o texto; admite coma o punto decimal (AR / US)."""
    if val is None or val == "":
        return None
    if isinstance(val, (int, float)):
        x = float(val)
        if x != x:  # NaN
            return None
        return _normalizar_puntaje_desde_hoja(x)
    s = unicodedata.normalize("NFKC", str(val))
    s = re.sub(r"\s+", "", s)
    if not s or s.lower() in ("nan", "none"):
        return None
    # Patrón claro: parte entera + un separador + parte decimal (evita ambigüedades)
    m = re.match(r"^(\d{1,4})([.,])(\d{1,4})$", s)
    if m:
        whole, _sep, frac = m.groups()
        try:
            return _normalizar_puntaje_desde_hoja(float(f"{whole}.{frac}"))
        except ValueError:
            return None
    # Miles con punto y decimal con coma: 1.234,56
    m2 = re.match(r"^(\d{1,3}(?:\.\d{3})+),(\d+)$", s)
    if m2:
        whole = m2.group(1).replace(".", "")
        frac = m2.group(2)
        try:
            return _normalizar_puntaje_desde_hoja(float(f"{whole}.{frac}"))
        except ValueError:
            return None
    # Solo dígitos (entero)
    m3 = re.match(r"^(\d{1,4})$", s)
    if m3:
        try:
            return _normalizar_puntaje_desde_hoja(float(m3.group(1)))
        except ValueError:
            return None
    # Fallback: un solo tipo de separador
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        s = s.replace(",", ".")
    try:
        return _normalizar_puntaje_desde_hoja(float(s))
    except ValueError:
        return None


def puntaje_a_texto_celda_sheet(x: float) -> str:
    """Texto con punto ASCII para la celda: evita que Sheets (locale ES) reinterprete comas."""
    x = float(x)
    if abs(x - round(x)) < 1e-9:
        return str(int(round(x)))
    return f"{x:.4f}".rstrip("0").rstrip(".")


def parse_anio_hoja(s: str) -> int | None:
    """Año como entero para la planilla (Looker Studio agrupa por número, no por texto)."""
    if s is None or not str(s).strip():
        return None
    s = str(s).strip()
    if not re.fullmatch(r"\d{4}", s):
        return None
    n = int(s)
    if n < 1990 or n > 2100:
        return None
    return n


def parse_puntaje_campo_formulario(s: str) -> tuple[float, str | None]:
    """Vacío → 0.0. Devuelve (valor, mensaje_error o None)."""
    if s is None or not str(s).strip():
        return 0.0, None
    n = parse_puntaje_valor(s)
    if n is None:
        return 0.0, "Use solo números; decimales con coma o punto (ej: 87,9 o 87.9)."
    if n < 0 or n > 1000:
        return 0.0, "El puntaje debe estar entre 0 y 1000."
    return n, None


def contar_palabras(texto: str) -> int:
    return len(re.findall(r"\S+", str(texto or "").strip()))


def format_puntaje_doc_es(x: float) -> str:
    """Texto para Word/correo: siempre 2 decimales y coma (ej. 86,00 como en la hoja)."""
    x = _normalizar_puntaje_desde_hoja(float(x))
    return f"{x:.2f}".replace(".", ",")


def puntaje_texto_para_word(raw) -> str | None:
    """
    Texto final para la línea «Puntaje: …» del Word.
    Refuerza la corrección x100 (hoja mal formateada / API) y evita mostrar 8780 tal cual.
    """
    if raw in (None, ""):
        return None
    n = parse_puntaje_valor(raw)
    if n is None or n <= 0:
        return None
    x = float(n)
    for _ in range(10):
        if x <= 1000:
            break
        # tolerancia por floats de la API (8780.000000001)
        if abs(x - round(x)) >= 1e-4:
            break
        ri = int(round(x))
        if ri % 100 != 0:
            break
        x = x / 100.0
    if x <= 0 or x > 1000:
        return None
    # Siempre dos decimales y coma decimal (alineado a Google Sheets: 86,00)
    return f"{x:.2f}".replace(".", ",")


def _container_con_estilo(key: str, border: bool = True):
    try:
        return st.container(border=border, key=key)
    except TypeError:
        return st.container(border=border)


def _texto_resumen(valor, vacio: str = "Pendiente") -> str:
    texto = str(valor or "").strip()
    return texto if texto else vacio


def _render_encabezado_bloque(paso: str, titulo: str, ayuda: str) -> None:
    st.markdown(
        f"""
        <p class="ucc-section-label">{escape(paso)}</p>
        <h3 class="ucc-section-title">{escape(titulo)}</h3>
        <p class="ucc-section-help">{escape(ayuda)}</p>
        """,
        unsafe_allow_html=True,
    )


def render_cabecera_carga_temas(numero_acta, fecha, tipo, cantidad_temas_acta: int | None) -> None:
    with _container_con_estilo("ucc_card_hero_carga_temas"):
        st.markdown(
            """
            <p class="ucc-section-label">Consejo de Investigación</p>
            <h1 class="ucc-section-title" style="font-size:2rem; margin-bottom:0.35rem;">
                Carga de temas al Orden del Día
            </h1>
            <p class="ucc-page-intro">
                Ingrese un tema, revise los datos clave y envíelo al acta correspondiente.
                La carga mantiene el flujo actual a Google Sheets y luego permite generar el Word.
            </p>
            """,
            unsafe_allow_html=True,
        )
        chips = [
            f'<span class="ucc-chip"><strong>Acta</strong> {_texto_resumen(numero_acta)}</span>',
            f'<span class="ucc-chip"><strong>Fecha</strong> {escape(_texto_resumen(fecha))}</span>',
            f'<span class="ucc-chip"><strong>Tipo</strong> {escape(_texto_resumen(tipo, "Seleccione un tema"))}</span>',
            f'<span class="ucc-chip"><strong>Temas cargados</strong> {_texto_resumen(cantidad_temas_acta, "Sin datos")}</span>',
        ]
        st.markdown(
            f'<div class="ucc-chip-row">{"".join(chips)}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            """
            <div class="ucc-card-note">
                Flujo sugerido: seleccione el acta, complete solo los campos que correspondan,
                revise el resumen y luego envíe. Después puede pasar a Carga de Archivos o generar el Orden del Día.
            </div>
            """,
            unsafe_allow_html=True,
        )
        col_a, col_b, col_c = st.columns(3)
        try:
            with col_a:
                st.page_link(
                    "pages/1_Cargar_Temas_al_Orden_del_Dia.py",
                    label="Cargar tema",
                    icon="📝",
                    use_container_width=True,
                )
            with col_b:
                st.page_link(
                    "pages/1_Descargar_Orden_del_Dia.py",
                    label="Generar OD",
                    icon="📄",
                    use_container_width=True,
                )
            with col_c:
                st.page_link(
                    "pages/2_Carga_de_Archivos.py",
                    label="Cargar archivo",
                    icon="📂",
                    use_container_width=True,
                )
        except AttributeError:
            with col_a:
                st.markdown('<a class="ucc-nav-btn-3d" href="/Cargar_Temas_al_Orden_del_Dia">Cargar tema</a>', unsafe_allow_html=True)
            with col_b:
                st.markdown('<a class="ucc-nav-btn-3d" href="/Descargar_Orden_del_Dia">Generar OD</a>', unsafe_allow_html=True)
            with col_c:
                st.markdown('<a class="ucc-nav-btn-3d" href="/Carga_de_Archivos">Cargar archivo</a>', unsafe_allow_html=True)


def render_resumen_pre_envio(
    numero_acta,
    fecha,
    anio,
    tipo,
    titulo,
    unidades_sel,
    responsable,
    director,
    descripcion,
) -> None:
    with _container_con_estilo("ucc_card_soft_resumen_envio"):
        _render_encabezado_bloque(
            "Revisión final",
            "Revisión antes de enviar",
            "Verifique los datos más importantes antes de confirmar la carga.",
        )
        resumen = [
            ("Acta", _texto_resumen(numero_acta)),
            ("Fecha", _texto_resumen(fecha)),
            ("Año", _texto_resumen(anio)),
            ("Tipo", _texto_resumen(tipo)),
            ("Título", _texto_resumen(titulo)),
            ("Unidad académica", _texto_resumen("; ".join(unidades_sel))),
            ("Responsable", _texto_resumen(responsable)),
            ("Director", _texto_resumen(director, "No corresponde")),
            ("Descripción", _texto_resumen(descripcion, "Sin descripción")),
        ]
        html_items = "".join(
            f'<div class="ucc-summary-item"><strong>{escape(label)}</strong><span>{escape(str(value))}</span></div>'
            for label, value in resumen
        )
        st.markdown(
            f'<div class="ucc-summary-grid">{html_items}</div>',
            unsafe_allow_html=True,
        )


def render_historial_acta(sheet, numero_acta) -> None:
    with _container_con_estilo("ucc_card_historial_acta"):
        _render_encabezado_bloque(
            "Control rápido",
            "Temas ya cargados en esta acta",
            "Vista de apoyo para evitar duplicados y revisar lo último ingresado.",
        )
        if not numero_acta:
            st.caption("Seleccione un acta para ver el historial inmediato.")
            return
        try:
            datos = sheet.get_all_records()
        except Exception as exc:
            st.caption(f"No se pudo leer el historial: {exc}")
            return
        registros = [
            _fila_sheet_normalizada(r)
            for r in datos
            if _numero_acta_igual(r.get("numero_acta"), numero_acta)
        ]
        registros = registros[-8:]
        if not registros:
            st.caption("Todavía no hay temas cargados para esta acta.")
            return
        for idx, r in enumerate(reversed(registros), start=1):
            titulo = str(r.get("titulo") or r.get("título") or "Sin título").strip()
            tipo = str(r.get("tipo") or "").strip() or "Sin tipo"
            unidad = str(r.get("unidad académica") or r.get("unidad") or "").strip() or "Sin unidad"
            responsable = str(r.get("responsable_de_carga") or "").strip() or "Sin responsable"
            st.markdown(
                f"""
                <div class="ucc-card-note" style="margin-top:{'0' if idx == 1 else '8px'};">
                    <strong>{idx}. {escape(tipo)}</strong><br>
                    <span>{escape(titulo)}</span><br>
                    <span style="color:#64748b;">{escape(unidad)} · {escape(responsable)}</span>
                </div>
                """,
                unsafe_allow_html=True,
            )


# =========================
# ⚙ CONFIGURACIÓN
# =========================

st.set_page_config(page_title="Consejo de Investigación", layout="wide")
hide_streamlit_cloud_toolbar()

_ir_a_descargar_od = st.session_state.pop("ir_a_descargar_orden_dia", False)
_viene_de_otra_pagina = st.session_state.get("_pagina_streamlit_prev") != "cargar_temas"
st.session_state["_pagina_streamlit_prev"] = "cargar_temas"

_scroll_arriba = (not _ir_a_descargar_od) and (
    st.session_state.pop("volver_arriba_cargar_temas", False) or _viene_de_otra_pagina
)

# Al elegir acta / generar, Streamlit re-ejecuta y vuelve arriba: mantener ancla en OD.
if _ir_a_descargar_od:
    st.session_state["_mantener_scroll_descargar_od"] = True
if _scroll_arriba:
    st.session_state.pop("_mantener_scroll_descargar_od", None)
_mantener_en_od = bool(st.session_state.get("_mantener_scroll_descargar_od"))

# Ancla superior: el menú «Cargar Temas» vuelve acá aunque ya estés en esta página.
st.markdown('<div id="inicio-cargar-temas"></div>', unsafe_allow_html=True)

components.html(
    f"""
    <script>
    (function () {{
        const win = window.parent;
        const doc = win.document;
        const storage = win.sessionStorage;
        const irOd = {"true" if _ir_a_descargar_od else "false"};
        const irArriba = {"true" if _scroll_arriba else "false"};

        function contenedoresScroll() {{
            return [
                win,
                doc.scrollingElement,
                doc.documentElement,
                doc.body,
                doc.querySelector("section.main"),
                doc.querySelector('[data-testid="stAppViewContainer"]'),
                doc.querySelector('[data-testid="stMainBlockContainer"]'),
                doc.querySelector('[data-testid="stMain"]'),
            ];
        }}

        function subir() {{
            storage.setItem("ucc_scroll_mode", "top");
            storage.removeItem("ucc_scroll_od");
            try {{ win.scrollTo(0, 0); }} catch (e) {{}}
            contenedoresScroll().forEach(function (el) {{
                if (!el) return;
                try {{
                    if (typeof el.scrollTo === "function") el.scrollTo(0, 0);
                    if ("scrollTop" in el) el.scrollTop = 0;
                }} catch (e) {{}}
            }});
            const ancla = doc.getElementById("inicio-cargar-temas");
            if (ancla) {{
                try {{ ancla.scrollIntoView({{ behavior: "auto", block: "start" }}); }} catch (e) {{}}
            }}
        }}

        function programarSubir() {{
            subir();
            [0, 50, 150, 350, 700, 1200, 2000].forEach(function (ms) {{
                setTimeout(subir, ms);
            }});
        }}

        if (!win._uccNavScrollInit) {{
            win._uccNavScrollInit = true;
            setInterval(function () {{
                if (storage.getItem("ucc_scroll_top") === "1") {{
                    storage.removeItem("ucc_scroll_top");
                    programarSubir();
                }}
            }}, 150);
            function enlazarMenu() {{
                doc.querySelectorAll(
                    '[data-testid="stSidebarNav"] a, [data-testid="stSidebarNavLink"], [data-testid="stSidebarNav"] li'
                ).forEach(function (a) {{
                    if (a._uccBound) return;
                    const t = (a.textContent || "").toLowerCase();
                    const esCargar = t.includes("cargar") && t.includes("temas");
                    const esDescargar = t.includes("descargar") && t.includes("orden");
                    if (!esCargar && !esDescargar) return;
                    a._uccBound = true;
                    a.addEventListener("pointerdown", function () {{
                        if (esCargar) {{
                            storage.setItem("ucc_scroll_mode", "top");
                            storage.setItem("ucc_scroll_top", "1");
                            storage.removeItem("ucc_scroll_od");
                            programarSubir();
                        }} else if (esDescargar) {{
                            storage.setItem("ucc_scroll_mode", "od");
                            storage.setItem("ucc_scroll_od", "1");
                            storage.removeItem("ucc_scroll_top");
                        }}
                    }}, true);
                    a.addEventListener("click", function () {{
                        if (esCargar) {{
                            storage.setItem("ucc_scroll_mode", "top");
                            storage.setItem("ucc_scroll_top", "1");
                            programarSubir();
                        }}
                    }}, true);
                }});
            }}
            enlazarMenu();
            new MutationObserver(enlazarMenu).observe(doc.body, {{
                childList: true,
                subtree: true,
            }});
        }}

        if (irOd) {{
            storage.setItem("ucc_scroll_mode", "od");
            storage.setItem("ucc_scroll_od", "1");
            storage.removeItem("ucc_scroll_top");
        }} else if (irArriba) {{
            storage.setItem("ucc_scroll_top", "1");
            programarSubir();
        }}
    }})();
    </script>
    """,
    height=0,
)

_APP_ROOT = Path(__file__).resolve().parent.parent
_LOGO_PATH = _APP_ROOT / "assets" / "logo_uccuyo.png"

# =========================
# 🎨 HEADER
# =========================

col1, col2 = st.columns([1, 8], vertical_alignment="center")

with col1:
    if _LOGO_PATH.is_file():
        st.image(str(_LOGO_PATH), width=120)
    else:
        st.caption("Logo no encontrado (assets/logo_uccuyo.png)")

with col2:
    st.markdown("""
    <div class='header-uccuyo'>
        <h2 style="font-weight:600;">Universidad Católica de Cuyo</h2>
        <h4 style="opacity:0.9;">Secretaría de Investigación</h4>
        <h5 style="opacity:0.8;">Consejo de Investigación</h5>
    </div>
    """, unsafe_allow_html=True)

# =========================
# 🎨 CSS GLOBAL
# =========================

st.markdown("""
<style>

/* Fondo principal (mint claro, estilo EvaluAR) */
.stApp,
[data-testid="stAppViewContainer"] {
    background-color: #D8EBE2 !important;
}

/* Sidebar (mint más marcado) */
[data-testid="stSidebar"] {
    background-color: #B5D5C6 !important;
}

[data-testid="stSidebar"] > div:first-child {
    background-color: #B5D5C6 !important;
}

section[data-testid="stSidebar"] * {
    color: #1e293b !important;
}

/* HEADER INSTITUCIONAL */
.header-uccuyo {
    background: linear-gradient(90deg, #064a3f, #0B6B5D);
    padding: 18px 20px;
    border-radius: 10px;
    display: flex;
    flex-direction: column;
    justify-content: center;
    box-shadow: 0 2px 8px rgba(6, 74, 63, 0.18);
}

.header-uccuyo h2,
.header-uccuyo h4,
.header-uccuyo h5 {
    color: white !important;
    margin: 0;
    font-weight: 600 !important;
}

/* Títulos generales: gris pizarra, no negro puro */
section.main h1,
section.main h2,
section.main h3,
section.main h4,
section.main h5,
section.main h6 {
    color: #1e293b !important;
    font-weight: 600 !important;
}

/* Labels Streamlit nativos */
label {
    color: #334155 !important;
    font-weight: 600 !important;
}

/* INPUTS */
input, textarea {
    background-color: white !important;
    color: #1e293b !important;
}

/* Selectbox */
div[data-baseweb="select"] {
    background-color: white !important;
}

div[data-baseweb="select"] > div {
    border-color: #cbd5e1 !important;
}

div[role="listbox"] {
    background-color: white !important;
}

div[role="option"] {
    background-color: white !important;
    color: #1e293b !important;
}

div[role="option"]:hover {
    background-color: #ecf5f2 !important;
}

/* Chips del multiselect (unidades académicas): verde UCCuyo, no rojo */
span[data-baseweb="tag"],
[data-baseweb="tag"] {
    background-color: #0B6B5D !important;
    color: #ffffff !important;
    border-radius: 6px !important;
    border: none !important;
}

span[data-baseweb="tag"] span,
[data-baseweb="tag"] span {
    color: #ffffff !important;
}

span[data-baseweb="tag"] svg,
[data-baseweb="tag"] svg {
    fill: #ffffff !important;
}

/* Barra vertical al final del texto en selects = caret del combobox Base Web */
div[data-baseweb="select"] input,
div[data-baseweb="select"] [data-baseweb="input"] input,
div[data-baseweb="select"] [role="combobox"],
div[data-baseweb="select"] [contenteditable="true"] {
    caret-color: transparent !important;
}

::placeholder {
    color: #94a3b8 !important;
}

.stTextInput > div > div > input {
    background-color: white !important;
    color: #1e293b !important;
    caret-color: #0B6B5D !important;
}

.stTextArea textarea {
    background-color: white !important;
    color: #1e293b !important;
    caret-color: #0B6B5D !important;
}

.stNumberInput input,
[data-testid="stNumberInputField"] {
    caret-color: #0B6B5D !important;
}

/* Sidebar: botones claros sobre mint (estilo EvaluAR) */
[data-testid="stSidebarNav"] a {
    background-color: #ffffff !important;
    color: #1e293b !important;
    font-size: 15px !important;
    margin-bottom: 6px;
    padding: 9px 12px;
    border-radius: 8px;
    font-weight: 500 !important;
    border: 1px solid #9bbbac !important;
    box-shadow: 0 1px 2px rgba(6, 74, 63, 0.06);
}

[data-testid="stSidebarNav"] a[aria-current="page"] {
    background-color: #9fc9b6 !important;
    color: #064a3f !important;
    font-weight: 600 !important;
    border-color: #9fc9b6 !important;
}

[data-testid="stSidebarNav"] a,
[data-testid="stSidebarNav"] a span,
[data-testid="stSidebarNav"] a p {
    color: inherit !important;
}

/* Alertas: menos énfasis */
[data-testid="stAlert"] * {
    color: #1e293b !important;
    font-weight: 500 !important;
}

/* Captions */
section.main [data-testid="stCaptionContainer"],
section.main [data-testid="stCaptionContainer"] p,
section.main [data-testid="stCaptionContainer"] small,
section.main [data-testid="stCaptionContainer"] span,
section.main [data-testid="stCaption"],
[data-testid="stMain"] [data-testid="stCaptionContainer"],
[data-testid="stMain"] [data-testid="stCaptionContainer"] p,
[data-testid="stMain"] [data-testid="stCaptionContainer"] span {
    color: #64748b !important;
    font-weight: 400 !important;
}

/* Estilo 3D bordó de botones: ucc_streamlit_chrome.inject_botones_3d_consejo() */

</style>
""", unsafe_allow_html=True)

# =========================
# 🔐 CONEXIÓN
# =========================

scope = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

creds = Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=scope
)

client = gspread.authorize(creds)

SHEET_ID = "17MiyW17W7oLIwSCKjDXCoA85CwBkYqHYhDKblVN37c8"
sheet = client.open_by_key(SHEET_ID).worksheet("Hoja 2")

st.markdown(
    """
    <div style='background-color:#ecf5f2; padding:12px 14px; border-radius:8px; border:1px solid #c5ddd6; margin-bottom:0.5rem;'>
        <a href='https://docs.google.com/spreadsheets/d/17MiyW17W7oLIwSCKjDXCoA85CwBkYqHYhDKblVN37c8' target='_blank' style='text-decoration:none; color:#0b6b5d; font-weight:600; font-size:15px;'>
            Base de datos de Órdenes del Día del Consejo de Investigación (abrir)
        </a>
    </div>
    """,
    unsafe_allow_html=True
)

# =========================
# 📅 DATOS ACTAS
# =========================

actas_dict = {
    187: {"mes": "Febrero"},
    188: {"mes": "Marzo"},
    189: {"mes": "Abril"},
    190: {"mes": "Mayo"},
    191: {"mes": "Junio"},
    192: {"mes": "Julio"},
    193: {"mes": "Agosto"},
    194: {"mes": "Septiembre"},
    195: {"mes": "Octubre"},
    196: {"mes": "Noviembre"},
    197: {"mes": "Diciembre"},
}

fechas_actas = {
    187: "19 de Febrero 2026",
    188: "19 de Marzo 2026",
    189: "16 de Abril 2026",
    190: "21 de Mayo 2026",
    191: "18 de Junio 2026",
    192: "23 de Julio 2026",
    193: "20 de Agosto 2026",
    194: "15 de Septiembre 2026",
    195: "22 de Octubre 2026",
    196: "19 de Noviembre 2026",
    197: "10 de Diciembre 2026"
}
# 🔹 LISTA DE CATEGORÍAS (VA ARRIBA DE TODO)
categoria_opciones = [
    "Seleccionar",
    "Investigador/a Superior I",
    "Investigador/a Principal II",
    "Investigador/a Independiente III",
    "Investigador/a Adjunto/a IV",
    "Investigador/a Asistente V",
    "Becario/a de Iniciación VI",
    "Sin categorización / Externo"
]

MAX_UNIDADES_ACADEMICAS = 5
opciones_unidades_select = [
    "FDCSSL- Facultad de Derecho y Ciencias Sociales Sede San Luis",
    "FCMSL- Facultad de Ciencias Médicas Sede San Luis",
    "FCVSL- Facultad de Ciencias Veterinarias Sede San Luis",
    "FCEESL- Facultad de Ciencias Económicas y Empresariales Sede San Luis",
    "FBOSCO- Facultad Don Bosco",
    "FCEESJ- Facultad de Ciencias Económicas y Empresariales Sede San Juan",
    "FFyHSJ- Facultad de Filosofía y Humanidades",
    "ISDSM- Instituto Universitario Santa María",
    "ECRyPSJ- Escuela Cultura Religiosa y Pastoral",
    "FDCSSJ- Facultad de Derecho y Ciencias Sociales Sede San Juan",
    "FCMSJ- Facultad de Ciencias Médicas San Juan",
    "FEDSJ- Facultad de Educación",
    "ESEGSJ- Escuela de Seguridad",
    "FCQyTSJ- Facultad de Ciencias Químicas y Tecnológicas",
    "ISB- Instituto San Buenaventura",
    "Secretaría de Investigación",
    "Unidad de Vinculación Tecnológica",
    "OIA- Observatorio de Inteligencia Artificial",
    "Vicerrectora de Formación",
    "Departamento de Educación a Distancia",
]

# =========================
# 📊 DASHBOARD estado (por acta y unidades)
# =========================

_EXPECTED_UNIDADES = set(opciones_unidades_select)
try:
    _datos_sheet_all = sheet.get_all_records()
except Exception:
    _datos_sheet_all = []

_por_acta_temas: dict[int, int] = {n: 0 for n in actas_dict.keys()}
_por_acta_unidades: dict[int, set[str]] = {n: set() for n in actas_dict.keys()}
_por_acta_unidad_counts: dict[int, dict[str, int]] = {n: {} for n in actas_dict.keys()}

for r in _datos_sheet_all:
    n_acta = _parse_numero_acta(r.get("numero_acta"))
    if n_acta is None or n_acta not in actas_dict:
        continue

    _por_acta_temas[n_acta] += 1

    unidad = _unidad_academica_clave(r)
    if unidad:
        _por_acta_unidades[n_acta].add(unidad)
        _por_acta_unidad_counts[n_acta][unidad] = (
            _por_acta_unidad_counts[n_acta].get(unidad, 0) + 1
        )

_TOTAL_EXPECTED_UNIDADES = len(_EXPECTED_UNIDADES) if _EXPECTED_UNIDADES else 1
_UMBRAL_COMPLETA = 0.85


def _estado_acta(numero_acta: int) -> tuple[str, str]:
    """Retorna (texto, background_css)."""
    temas = _por_acta_temas.get(numero_acta, 0)
    if temas <= 0:
        return "Sin carga", "background:#b91c1c"

    unidades_presentes = len(_por_acta_unidades[numero_acta] & _EXPECTED_UNIDADES)
    cobertura = unidades_presentes / _TOTAL_EXPECTED_UNIDADES
    if cobertura >= _UMBRAL_COMPLETA:
        return "Completa", "background:#0b6b5d"
    return "Parcial", "background:#b45309"


st.markdown("## 📊 Dashboard de estado")
st.caption("Señala cuántos temas hay y qué cobertura de unidades ya tiene cada Acta.")

st.markdown(
    """
    <style>
    .ucc-badge-pill{
      display:inline-flex;
      align-items:center;
      justify-content:center;
      padding: 4px 10px;
      border-radius: 999px;
      color: #fff;
      font-weight: 800;
      font-size: 0.85rem;
      line-height: 1;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

_actas_ordenadas = sorted(actas_dict.keys())
cols_actas = st.columns(4)
for idx, n in enumerate(_actas_ordenadas):
    c = cols_actas[idx % 4]
    temas = _por_acta_temas.get(n, 0)
    unidades_presentes = len(_por_acta_unidades[n] & _EXPECTED_UNIDADES)
    estado_txt, estado_bg = _estado_acta(n)

    mes = actas_dict[n]["mes"]
    with c:
        st.markdown(
            f"""
            <div class="ucc-summary-item" style="padding:14px 12px;">
              <div style="display:flex; align-items:center; justify-content:space-between; gap:10px;">
                <strong style="color:#334155;">Acta {n}</strong>
                <span class="ucc-badge-pill" style="{estado_bg}">{estado_txt}</span>
              </div>
              <div style="margin-top:4px; color:#064a3f; font-weight:600; font-size:0.92rem;">
                {mes}
              </div>
              <div style="margin-top:6px; color:#475569; font-weight:700;">
                {temas} tema(s)
              </div>
              <div style="margin-top:4px; color:#64748b;">
                {unidades_presentes}/{_TOTAL_EXPECTED_UNIDADES} unidades con carga
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

# =========================
# 📝 FORMULARIO
# =========================

OPCION_ACTA_SIN_SELECCION = "Seleccionar el Orden del día"
opciones_acta_carga = [OPCION_ACTA_SIN_SELECCION] + [
    f"Orden del Día {actas_dict[n]['mes']} - Acta {n}" for n in actas_dict
]

acta_label = st.session_state.get("acta", OPCION_ACTA_SIN_SELECCION)
if acta_label == OPCION_ACTA_SIN_SELECCION:
    numero_acta = None
    fecha = ""
else:
    numero_acta = int(acta_label.split("Acta ")[1])
    fecha = fechas_actas.get(numero_acta, "")

cantidad_temas_acta = _por_acta_temas.get(numero_acta, 0) if numero_acta else None
tipo_actual = st.session_state.get("tipo_actividad", "Proyecto de Investigación")
render_cabecera_carga_temas(numero_acta, fecha, tipo_actual, cantidad_temas_acta)

col_kpi_1, col_kpi_2, col_kpi_3 = st.columns(3)
with col_kpi_1:
    with _container_con_estilo("ucc_card_kpi_acta"):
        st.markdown(
            f"""
            <p class="ucc-kpi-label">Acta en trabajo</p>
            <p class="ucc-kpi-value">{escape(_texto_resumen(numero_acta, "Sin seleccionar"))}</p>
            <p class="ucc-kpi-sub">Seleccione primero el Orden del Día para asociar correctamente el tema.</p>
            """,
            unsafe_allow_html=True,
        )
with col_kpi_2:
    with _container_con_estilo("ucc_card_kpi_temas"):
        st.markdown(
            f"""
            <p class="ucc-kpi-label">Temas cargados</p>
            <p class="ucc-kpi-value">{escape(_texto_resumen(cantidad_temas_acta, "0"))}</p>
            <p class="ucc-kpi-sub">Conteo actual de registros asociados al acta seleccionada.</p>
            """,
            unsafe_allow_html=True,
        )
with col_kpi_3:
    with _container_con_estilo("ucc_card_kpi_accion"):
        st.markdown(
            f"""
            <p class="ucc-kpi-label">Siguiente acción</p>
            <p class="ucc-kpi-value">{escape("Completar carga" if numero_acta else "Elegir acta")}</p>
            <p class="ucc-kpi-sub">Mantenga la lógica actual: cargar, revisar y luego generar el Word del OD.</p>
            """,
            unsafe_allow_html=True,
        )

if numero_acta:
    _unidades_counts = _por_acta_unidad_counts.get(numero_acta, {})
    _unidades_presentes = set(_unidades_counts.keys())
    _unidades_top = sorted(_unidades_counts.items(), key=lambda kv: kv[1], reverse=True)[:6]
    _unidades_pendientes = sorted(_EXPECTED_UNIDADES - _unidades_presentes)

    with _container_con_estilo("ucc_card_unidades_acta"):
        st.markdown("### Unidades con temas (Acta seleccionado)")
        st.caption(
            f"Top unidades con carga y cuántas unidades faltan completar (según el listado del sistema)."
        )

        st.markdown(
            "<div class='ucc-chip-row'>"
            + "".join(
                f"<span class='ucc-chip'><strong>{escape(str(c[1]))}</strong>&nbsp;{escape(str(c[0]))}</span>"
                for c in _unidades_top
            )
            + "</div>",
            unsafe_allow_html=True,
        )

        faltan_n = len(_unidades_pendientes)
        st.markdown(
            f"<div style='margin-top:10px; color:#64748b; font-weight:700;'>Unidades pendientes: {faltan_n}</div>",
            unsafe_allow_html=True,
        )
        if faltan_n > 0:
            with st.expander("Ver unidades pendientes (detalle)"):
                st.markdown(
                    "<br>".join(
                        escape(u) for u in _unidades_pendientes[:80]
                    )
                    if _unidades_pendientes
                    else "N/A",
                    unsafe_allow_html=False,
                )

with _container_con_estilo("ucc_card_paso_1"):
    _render_encabezado_bloque(
        "Paso 1",
        "Datos generales del acta",
        "Seleccione el año, el Orden del Día y confirme la fecha de reunión antes de continuar.",
    )
    col_bas_1, col_bas_2, col_bas_3 = st.columns([1, 2, 2])

    with col_bas_1:
        st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Año</div>", unsafe_allow_html=True)
        anio = st.text_input("", "2026", key="anio")

    with col_bas_2:
        st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Seleccione el Orden del Día</div>", unsafe_allow_html=True)
        acta_label = st.selectbox(
            "",
            opciones_acta_carga,
            index=0,
            key="acta",
        )

    if acta_label == OPCION_ACTA_SIN_SELECCION:
        numero_acta = None
        fecha = ""
    else:
        numero_acta = int(acta_label.split("Acta ")[1])
        fecha = fechas_actas.get(numero_acta, "")

    with col_bas_3:
        st.markdown("<div style='margin-bottom:6px; color:#334155; font-weight:600; font-size:0.95rem;'>Fecha de la reunión de Consejo de Investigación</div>", unsafe_allow_html=True)
        st.markdown(
            f"<p style='color:#475569; margin:0 0 20px 0; padding-bottom:4px; font-weight:500;'>{escape(fecha)}</p>",
            unsafe_allow_html=True,
        )

with _container_con_estilo("ucc_card_paso_2"):
    _render_encabezado_bloque(
        "Paso 2",
        "Tipo de actividad o tema",
        "El sistema mostrará más abajo solo los campos que correspondan al tipo elegido.",
    )
    col_tipo_1, col_tipo_2 = st.columns([2, 3], vertical_alignment="top")
    with col_tipo_1:
        st.markdown(
            "<div style='margin:0 0 10px 0; color:#334155; font-weight:600; font-size:0.95rem; line-height:1.3;'>Actividad o Tema</div>",
            unsafe_allow_html=True,
        )
        tipo = st.selectbox(
            "Tipo de actividad",
            [
                "Proyecto de Investigación",
                "Proyecto de Cátedra",
                "Informe Final",
                "Informe de Avance",
                "Jornada de Investigación",
                "Convocatoria a Proyectos de investigación",
                "Creación de Semillero de Investigación",
                "Categorización Docente",
                "Llamado a Concurso de Becas",
                "Líneas prioritarias de investigación",
                "Otra",
            ],
            key="tipo_actividad",
            label_visibility="collapsed",
        )

    catedra_lateral = ""
    with col_tipo_2:
        if tipo == "Proyecto de Cátedra":
            st.markdown(
                "<div style='margin:0 0 10px 0; color:#334155; font-weight:600; font-size:0.95rem; line-height:1.3;'>"
                "Cátedra o cátedras <span style='color:#94a3b8;font-weight:500;'>(escribir a mano)</span>"
                "</div>",
                unsafe_allow_html=True,
            )
            catedra_lateral = st.text_input(
                "",
                key="catedra_proyecto_catedra",
                placeholder="Ej: Anatomía I; Fisiología; Clínica Médica",
                label_visibility="collapsed",
            )

with st.form("form_acta", clear_on_submit=False):

    # =========================
    # 📌 IDENTIFICACIÓN
    # =========================
    with _container_con_estilo("ucc_card_paso_3"):
        _render_encabezado_bloque(
            "Paso 3",
            "Denominación y descripción",
            "Cargue el nombre del tema y una descripción breve. Si corresponde, agregue puntaje.",
        )

        # Mockup: 3 columnas — Denominación+input | Indicaciones (solo banner) | Puntaje+input
        _hdr_iframe_h = 100
        col_den, col_ind, col_pun = st.columns([2.5, 2.5, 1.05], vertical_alignment="top")

        with col_den:
            _ayuda_en_iframe(
                f"<div style=\"box-sizing:border-box;height:{_hdr_iframe_h - 2}px;display:flex;align-items:center;"
                "padding:8px 12px;font:600 14px/1.25 system-ui,sans-serif;color:#334155;background:#f1f5f9;"
                'border-radius:6px;border:1px solid #e2e8f0;">Denominación de la actividad o Tema</div>',
                alto=_hdr_iframe_h,
            )

        with col_ind:
            _ayuda_en_iframe(
                f"<div style=\"box-sizing:border-box;min-height:{_hdr_iframe_h - 2}px;padding:8px 10px;font:11.5px/1.35 system-ui,sans-serif;"
                "color:#475569;background:#f1f5f9;border-radius:6px;border-left:4px solid #0b6b5d;"
                'border-top:1px solid #e2e8f0;border-right:1px solid #e2e8f0;border-bottom:1px solid #e2e8f0;">'
                "<strong>Indicaciones:</strong> "
                "Título del proyecto; Título del Informe Final o de Avance; "
                "Título de Jornada / Semillero / Instituto u otra actividad</div>",
                alto=_hdr_iframe_h,
            )

        with col_pun:
            if tipo in TIPOS_CON_PUNTAJE:
                _ayuda_en_iframe(
                    f"<div style=\"box-sizing:border-box;height:{_hdr_iframe_h - 2}px;display:flex;flex-direction:column;"
                    "justify-content:center;gap:6px;padding:8px 10px;font-family:system-ui,sans-serif;"
                    "background:#f1f5f9;border-radius:6px;border:1px solid #e2e8f0;\">"
                    "<div style=\"font-weight:600;color:#334155;font-size:14px;line-height:1.15;\">Puntaje</div>"
                    "<div style=\"font-size:11.5px;line-height:1.35;color:#64748b;\">"
                    "Decimales con coma o punto (ej: 87,9 o 87.9).</div></div>",
                    alto=_hdr_iframe_h,
                )

        # Inputs: denominación ancha (hasta borde de Indicaciones) | puntaje a la derecha
        col_tit_w, col_pun_inp = st.columns([5.0, 1.05], vertical_alignment="top")
        with col_tit_w:
            titulo = st.text_input("", key="titulo_actividad_consejo")
        with col_pun_inp:
            puntaje = 0.0
            if tipo in TIPOS_CON_PUNTAJE:
                puntaje_raw = st.text_input(
                    "",
                    placeholder="Ej: 87,9",
                    key="puntaje_informe_consejo",
                    label_visibility="collapsed",
                )
                _pv = parse_puntaje_valor(puntaje_raw)
                puntaje = _pv if _pv is not None else 0.0

        # Descripción solo bajo Denominación + Indicaciones (misma proporción que arriba: 2.5+2.5 vs 1.05)
        col_desc_w, _col_desc_pad = st.columns([5.0, 1.05])
        with col_desc_w:
            st.markdown(
                "<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Descripción (no más de 50 palabras)</div>",
                unsafe_allow_html=True,
            )
            descripcion = st.text_area("", key="descripcion")

    # =========================
    # 👥 EQUIPO (CONDICIONAL)
    # =========================
    with _container_con_estilo("ucc_card_paso_4"):
        _render_encabezado_bloque(
            "Paso 4",
            "Participantes y datos específicos",
            "Según el tipo elegido, complete docente, dirección, equipo, instituto, cátedra y alumnos.",
        )

        director = ""
        cat_director = ""
        codirector = ""
        categoria_codirector = ""
        equipo = ""
        instituto = ""
        catedra = ""
        alumnos = ""
        apellido_nombre_docente = ""
        dni_docente = ""

        if tipo == "Categorización Docente":
            col_doc_1, col_doc_2 = st.columns(2)
            with col_doc_1:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Apellido y nombre del docente</div>", unsafe_allow_html=True)
                apellido_nombre_docente = st.text_input("", key="apellido_nombre_docente")
            with col_doc_2:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>DNI</div>", unsafe_allow_html=True)
                dni_docente = st.text_input("", key="dni_docente")

        if tipo in ["Proyecto de Investigación", "Proyecto de Cátedra", "Informe Final", "Informe de Avance", "Otra"]:

            col_dir_1, col_dir_2 = st.columns(2)
            with col_dir_1:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Director</div>", unsafe_allow_html=True)
                director = st.text_input("", key="director")
            with col_dir_2:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Categoría del Director</div>", unsafe_allow_html=True)
                cat_director = st.selectbox("", categoria_opciones, key="cat_director")

            col_codir_1, col_codir_2 = st.columns(2)
            with col_codir_1:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Codirector</div>", unsafe_allow_html=True)
                codirector = st.text_input("", key="codirector")
            with col_codir_2:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Categoría del Codirector</div>", unsafe_allow_html=True)
                categoria_codirector = st.selectbox("", categoria_opciones, key="cat_codirector")

            st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Equipo de Investigación (no más de 50 palabras)</div>", unsafe_allow_html=True)
            equipo = st.text_area("", key="equipo", height=160)

            if tipo == "Proyecto de Cátedra":
                # La cátedra se carga al lado del tipo (arriba); acá solo instituto / alumnos.
                catedra = catedra_lateral
                col_eq_1, col_eq_2 = st.columns(2)
                with col_eq_1:
                    st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Instituto de Investigación</div>", unsafe_allow_html=True)
                    instituto = st.text_input("", key="instituto")
                with col_eq_2:
                    st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Número de Alumnos en el proyecto</div>", unsafe_allow_html=True)
                    alumnos = st.text_input("", key="alumnos")
            else:
                col_eq_1, col_eq_2, col_eq_3 = st.columns(3)
                with col_eq_1:
                    st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Instituto de Investigación</div>", unsafe_allow_html=True)
                    instituto = st.text_input("", key="instituto")
                with col_eq_2:
                    st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Cátedra (Si corresponde)</div>", unsafe_allow_html=True)
                    catedra = st.text_input("", key="catedra")
                with col_eq_3:
                    st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem; font-size:0.92rem; line-height:1.2;'>Número de Alumnos en el proyecto</div>", unsafe_allow_html=True)
                    alumnos = st.text_input("", key="alumnos")

    # =========================
    # 🏫 UNIDAD
    # =========================

    with _container_con_estilo("ucc_card_paso_45"):
        _render_encabezado_bloque(
            "Paso 5",
            "Unidad académica, resoluciones y financiamiento",
            "Asocie correctamente el tema a la unidad, complete resoluciones si corresponde y agregue el responsable de carga.",
        )

        col_uni_res_1, col_uni_res_2, col_uni_res_3 = st.columns([2.9, 1.05, 1.05])
        with col_uni_res_1:
            st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Unidad Académica</div>", unsafe_allow_html=True)
            st.caption(
                f"Máximo {MAX_UNIDADES_ACADEMICAS} unidades. "
                f"Con {MAX_UNIDADES_ACADEMICAS} elegidas, quite una con la × para cambiar."
            )
            unidades_sel = st.multiselect(
                "",
                opciones_unidades_select,
                key="unidades_academicas",
                max_selections=MAX_UNIDADES_ACADEMICAS,
                label_visibility="collapsed",
            )

        if tipo in ["Proyecto de Investigación", "Proyecto de Cátedra", "Informe Final", "Informe de Avance", "Otra"]:

            with col_uni_res_2:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Resolución CD</div>", unsafe_allow_html=True)
                resolucion_cd = st.text_input("", key="resolucion_cd", max_chars=10, placeholder="Ej: 665")
            with col_uni_res_3:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Resolución CS</div>", unsafe_allow_html=True)
                resolucion_cs = st.text_input("", key="resolucion_cs", max_chars=10, placeholder="Ej: 657")

        else:
            resolucion_cd = ""
            resolucion_cs = ""

        col_fin_1, col_fin_2, col_fin_3, col_fin_4 = st.columns(4)
        with col_fin_1:
            st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Responsable de carga <span style='color:#94a3b8;font-weight:500;'>(requerido)</span></div>", unsafe_allow_html=True)
            responsable_de_carga = st.text_input("", key="responsable")

        if tipo != "Categorización Docente":

            with col_fin_2:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Tipo de financiamiento</div>", unsafe_allow_html=True)
                tipo_financiamiento = st.selectbox("", ["Seleccionar...", "Interno", "Externo", "Sin financiamiento"], key="fin")
            with col_fin_3:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Fuente de Financiamiento</div>", unsafe_allow_html=True)
                fuente_financiamiento = st.text_input("", key="fuente")
            with col_fin_4:
                st.markdown("<div style='margin-bottom:-8px; color:#334155; font-weight:600; font-size:0.95rem;'>Monto en pesos (sin puntos)</div>", unsafe_allow_html=True)
                monto_financiamiento = st.number_input("", min_value=0, step=1000, value=None, key="monto")

        else:
            tipo_financiamiento = ""
            fuente_financiamiento = ""
            monto_financiamiento = 0

    # =========================
    # 🔘 SUBMIT
    # =========================

    render_resumen_pre_envio(
        numero_acta,
        fecha,
        anio,
        tipo,
        st.session_state.get("titulo_actividad_consejo", ""),
        st.session_state.get("unidades_academicas", []),
        st.session_state.get("responsable", ""),
        st.session_state.get("director", ""),
        st.session_state.get("descripcion", ""),
    )

    submit = st.form_submit_button("Clic para enviar al Consejo de Investigación (Google Sheets)")
    st.markdown(
        """
        <style>
        div[data-testid="stFormSubmitButton"] button {
            font-size: 18px !important;
            font-weight: 600 !important;
            border-radius: 10px !important;
            min-height: 3.2rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

def enviar_correo_tema(fila):

    destinatarios = [
        "investigacion@uccuyo.edu.ar",
        "vincutec@uccuyo.edu.ar",
        "asistente.inv@uccuyo.edu.ar"
    ]

    cuerpo = f"""
Se ha cargado un nuevo tema para el Consejo de Investigación.

Número de Acta: {fila[0]}
Fecha: {fila[1]}
Año: {fila[2]}
Tipo: {fila[3]}
Título: {fila[4]}
Descripción: {fila[5]}
Director: {fila[6]}
Codirector: {fila[8]}
Equipo: {fila[10]}
Unidad Académica: {fila[13]}
Resolución CD: {fila[14]}
Resolución CS: {fila[15]}
Instituto: {fila[16]}
Cátedra: {fila[17]}
Financiamiento: {fila[18]}
Fuente: {fila[19]}
Monto: {fila[20]}
Alumnos: {fila[21]}
Puntaje: {puntaje_texto_para_word(fila[22]) or "N/D"}
Responsable de carga: {fila[23]}
"""

    msg = EmailMessage()
    msg["Subject"] = "Nuevo tema cargado - Consejo de Investigación"
    msg["From"] = st.secrets["email"]["EMAIL_USER"]
    msg["To"] = ", ".join(destinatarios)
    msg.set_content(cuerpo)
    
    context = ssl.create_default_context()

    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls(context=context)

        user = str(st.secrets["email"]["EMAIL_USER"]).strip()
        # Gmail muestra la contraseña de aplicación con espacios; SMTP no los admite.
        password = str(st.secrets["email"]["EMAIL_PASS"]).replace(" ", "").strip()
        server.login(user, password)

        server.send_message(msg)

# =========================
# 💾 GUARDAR
# =========================

if submit:

    if tipo == "Proyecto de Cátedra":
        catedra = st.session_state.get("catedra_proyecto_catedra", catedra)

    if tipo in TIPOS_CON_PUNTAJE:
        puntaje_fila, err_puntaje = parse_puntaje_campo_formulario(
            st.session_state.get("puntaje_informe_consejo", "")
        )
    else:
        puntaje_fila = 0.0
        err_puntaje = None

    if tipo_financiamiento == "Seleccionar...":
        tipo_financiamiento = ""

    unidad = "; ".join(unidades_sel)

    # 🔹 LIMPIAR "Seleccionar" (versión robusta)
    if instituto.strip().startswith("Seleccionar"):
        instituto = ""

    catedra = str(catedra or "").strip()
    if catedra.startswith("Seleccionar"):
        catedra = ""
    else:
        # Una o varias cátedras en una sola celda / línea del Word
        catedra = "; ".join(
            parte.strip()
            for parte in catedra.replace("\n", ";").split(";")
            if parte.strip()
        )

    if tipo_financiamiento.strip().startswith("Seleccionar"):
        tipo_financiamiento = ""

    if str(cat_director).strip().startswith("Seleccionar"):
        cat_director = ""

    if str(categoria_codirector).strip().startswith("Seleccionar"):
        categoria_codirector = ""

    if monto_financiamiento is None:
        monto_financiamiento = ""

    anio_hoja = parse_anio_hoja(anio)

    fila = [
        int(numero_acta) if numero_acta is not None else numero_acta,
        fecha,
        anio_hoja,
        tipo,
        titulo,
        descripcion,
        director,
        cat_director,
        codirector,
        categoria_codirector,
        equipo,
        apellido_nombre_docente,
        dni_docente,
        unidad,
        resolucion_cd,
        resolucion_cs,
        instituto,
        catedra,
        tipo_financiamiento,
        fuente_financiamiento,
        monto_financiamiento,
        alumnos,
        puntaje_a_texto_celda_sheet(puntaje_fila),
        responsable_de_carga
        ]

    # VALIDACIONES
    if anio_hoja is None:
        st.error("Debe ingresar un año válido de cuatro dígitos (ej: 2026)")

    elif not numero_acta:
        st.error("Debe seleccionar el Orden del día")

    elif not fecha:
        st.error("Debe seleccionar la fecha")

    elif not tipo:
        st.error("Debe elegir la actividad")

    elif not titulo.strip():
        st.error("Debe completar la Denominación de la actividad")

    elif not unidades_sel:
        st.error("Debe seleccionar al menos una Unidad Académica")

    elif len(unidades_sel) > MAX_UNIDADES_ACADEMICAS:
        st.error(
            f"Solo puede seleccionar hasta {MAX_UNIDADES_ACADEMICAS} unidades académicas. "
            "Quite una con la × e intente de nuevo."
        )

    elif contar_palabras(descripcion) > 50:
        st.error("La descripción no debe superar 50 palabras")

    elif contar_palabras(equipo) > 50:
        st.error("El equipo de investigación no debe superar 50 palabras")

    elif not responsable_de_carga.strip():
        st.error("Debe completar el Responsable de carga")

    elif tipo == "Proyecto de Cátedra" and not str(catedra).strip():
        st.error("Debe indicar la cátedra o las cátedras del proyecto")

    elif err_puntaje:
        st.error(err_puntaje)

    else:
        sheet.append_row(
            fila,
            value_input_option=ValueInputOption.user_entered,
            insert_data_option=InsertDataOption.insert_rows,
        )

        try:
            enviar_correo_tema(fila)
        except Exception:
            # El guardado en Sheets ya se hizo. No mostrar aviso a quien carga:
            # Gmail (EMAIL_PASS) hay que renovarlo en secretos de Streamlit Cloud.
            pass

        n_acta = _contar_temas_acta(sheet, numero_acta)
        st.success(
            "✅ Registro guardado correctamente (se agregó al Orden del Día; no reemplaza temas anteriores).\n\n"
            f"El Acta {numero_acta} tiene ahora **{n_acta}** tema(s) cargado(s).\n\n"
            "Puede enviar otro tema ahora mismo (cambie los campos y vuelva a enviar) "
            "o cargar el archivo correspondiente."
        )
        st.markdown("""
        <a href="/Carga_de_Archivos" target="_self">
            <button class="ucc-btn-3d" style="
                width:auto;
                display:inline-block;
                text-align:left;
                padding:12px 20px;
                font-size:17px;
                margin-top:10px;">
                📂 Ahora cargue el archivo correspondiente
            </button>
        </a>
        """, unsafe_allow_html=True)
        
        st.markdown("<br><br>", unsafe_allow_html=True)
        
        
    
# =========================
# 📄 GENERAR WORD
# =========================

render_historial_acta(sheet, numero_acta)

st.markdown('<div id="descargar-orden-del-dia"></div>', unsafe_allow_html=True)

if _mantener_en_od:
    _scroll_od_suave = "true" if _ir_a_descargar_od else "false"
    components.html(
        f"""
        <script>
        (function () {{
            const win = window.parent;
            const doc = win.document;
            const storage = win.sessionStorage;
            const suave = {_scroll_od_suave};
            storage.setItem("ucc_scroll_mode", "od");
            storage.setItem("ucc_scroll_od", "1");

            function bajar() {{
                // Si el usuario pidió volver arriba, no seguir bajando.
                if (storage.getItem("ucc_scroll_mode") === "top") return true;
                if (storage.getItem("ucc_scroll_top") === "1") return true;
                const el = doc.getElementById("descargar-orden-del-dia");
                if (!el) return false;
                el.scrollIntoView({{
                    behavior: suave ? "smooth" : "auto",
                    block: "start"
                }});
                return true;
            }}
            // Primera visita (menú): reintentos largos. Re-runs (selectbox): inmediato.
            const delays = suave
                ? [0, 300, 700, 1200, 2000, 3500]
                : [0, 50, 150, 400, 800];
            delays.forEach(function (ms) {{
                setTimeout(bajar, ms);
            }});
        }})();
        </script>
        """,
        height=0,
    )

with _container_con_estilo("ucc_card_od_modulo"):
    _render_encabezado_bloque(
        "Módulo de gestión",
        "Generar y descargar Orden del Día",
        "Seleccione el acta, revise cuántos temas tiene cargados y ordene la salida antes de generar el Word.",
    )

    OPCION_OD_SIN_SELECCION = "Seleccione el orden del día"
    opciones_od_word = [OPCION_OD_SIN_SELECCION] + [
        f"{n} - {actas_dict[n]['mes']}" for n in actas_dict
    ]

    acta_word = st.selectbox(
        "Seleccionar Orden del Día para generar y descargar",
        options=opciones_od_word,
        index=0,
        key="acta_word_descargar",
    )

    registros_od = []
    acta_num_od = None
    if acta_word != OPCION_OD_SIN_SELECCION:
        acta_num_od = int(acta_word.split(" - ")[0])
        try:
            _datos_od = sheet.get_all_records()
        except Exception as exc:
            st.error(f"No se pudo leer la planilla: {exc}")
            _datos_od = []
        _base_od = [
            r for r in _datos_od
            if _numero_acta_igual(r.get("numero_acta"), acta_num_od)
        ]
        registros_od = _aplicar_orden_manual_od(acta_num_od, _base_od)
        if registros_od:
            st.info(f"Se encontraron **{len(registros_od)}** tema(s) para el Acta {acta_num_od}.")

            try:
                _caja_od = st.container(border=True, key="od_acciones_bordo")
            except TypeError:
                _caja_od = st.container(border=True)
            with _caja_od:
                col_rest, col_gen, col_dl = st.columns(3)
                with col_rest:
                    if st.button(
                        "Restaurar orden de carga",
                        key=f"od_reset_orden_{acta_num_od}",
                        use_container_width=True,
                        type="primary",
                    ):
                        st.session_state.pop(_clave_orden_manual_od(acta_num_od), None)
                        st.session_state.pop(f"od_docx_bytes_{acta_num_od}", None)
                        st.rerun()
                with col_gen:
                    generar = st.button(
                        "Generar Orden del Día",
                        key=f"od_generar_{acta_num_od}",
                        use_container_width=True,
                        type="primary",
                    )
                with col_dl:
                    _docx = st.session_state.get(f"od_docx_bytes_{acta_num_od}")
                    if _docx:
                        st.download_button(
                            "Descargar Orden del Día",
                            data=_docx,
                            file_name=f"Acta_{acta_num_od}.docx",
                            key=f"od_descargar_{acta_num_od}",
                            use_container_width=True,
                        )
                    else:
                        st.button(
                            "Descargar Orden del Día",
                            key=f"od_descargar_disabled_{acta_num_od}",
                            use_container_width=True,
                            disabled=True,
                            type="primary",
                            help="Primero genere el Orden del Día",
                        )

            if generar:
                st.session_state[f"od_docx_bytes_{acta_num_od}"] = _construir_bytes_orden_del_dia(
                    acta_num_od, registros_od
                )
                st.session_state[f"od_docx_ok_{acta_num_od}"] = True
                st.rerun()

            if st.session_state.pop(f"od_docx_ok_{acta_num_od}", False):
                st.success(
                    f"Orden del Día del Acta {acta_num_od} generado. "
                    "Use **Descargar Orden del Día** en la misma fila."
                )

            _ui_reordenar_temas_od(acta_num_od, registros_od)
            registros_od = _aplicar_orden_manual_od(acta_num_od, _base_od)
        else:
            st.caption("No hay temas cargados para esta acta todavía.")
            generar = False
    else:
        generar = False
        
with _container_con_estilo("ucc_card_soft_informe_responsable"):
    _render_encabezado_bloque(
        "Consulta rápida",
        "Generar informe por responsable",
        "Obtenga un Word con los temas cargados por una persona responsable dentro del acta seleccionada.",
    )

    responsable_reporte = st.text_input("Responsable de carga para generar informe")

    generar_responsables = st.button("Generar informe del responsable de carga")

    if generar_responsables:

        if acta_word == OPCION_OD_SIN_SELECCION:
            st.warning("Seleccione un orden del día antes de generar el informe.")
        else:
            datos = sheet.get_all_records()

            acta_num = int(acta_word.split(" - ")[0])

            registros = [
                r for r in datos
                if _numero_acta_igual(r.get("numero_acta"), acta_num)
                and str(r.get("responsable_de_carga", "")).strip().lower() == responsable_reporte.strip().lower()
            ]
            registros = ordenar_registros_por_unidad_academica(registros)

            if not responsable_reporte.strip():
                st.warning("Debe ingresar el responsable de carga")

            elif not registros:
                st.warning("No hay registros cargados por ese responsable para esta acta")

            else:
                doc = Document()

                doc.add_heading("Informe de temas cargados", 0)
                doc.add_paragraph(f"Acta N° {acta_num}")
                doc.add_paragraph(f"Responsable de carga: {responsable_reporte}")

                contador = 1
                unidad_actual = None

                for r in registros:

                    r = {k.lower().strip(): v for k, v in r.items()}

                    unidad = r.get("unidad académica", r.get("unidad", "")).strip()

                    if unidad != unidad_actual:
                        h = doc.add_paragraph()
                        h.paragraph_format.space_before = Pt(6)
                        h.paragraph_format.space_after = Pt(2)

                        run_h = h.add_run(unidad)
                        run_h.bold = True
                        run_h.font.color.rgb = RGBColor(0, 102, 204)

                        unidad_actual = unidad

                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1

                    p.add_run(f"{contador}. {r.get('tipo', '')} - {r.get('titulo', '')}\n").bold = True

                    descripcion = r.get("descripcion") or r.get("descripción") or ""

                    if descripcion:
                        p.add_run(f"   Descripción: {descripcion}\n")

                    p.add_run(f"   Unidad Académica: {unidad}\n")

                    contador += 1

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    "Descargar informe del responsable de carga",
                    data=buffer,
                    file_name=f"Informe_{responsable_reporte}_Acta_{acta_num}.docx"
                )
